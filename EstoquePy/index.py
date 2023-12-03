import tkinter as tk
from tkinter import messagebox
import pandas as pd
import os
from docx import Document
from docx.shared import Inches

# adicionando a função de salvar em .doc
def save_data():
    global data_df
    data_df = pd.DataFrame(list_data, columns=["Name ", "Quantity"])
    data_df.to_excel('stock_control.xlsx', index=False)

    doc = Document()
    doc.add_heading('stock control system', 0)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Product Name'
    header_cells[0].text = 'Quantity '

    for data in list_data:
        row_cells = table.add_row().cells
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
    
    doc.save('stock_control.docx')
    messagebox.showinfo('Salvo com sucesso')

def save_data():
    global data_df
    data_df = pd.DataFrame(list_data, columns=["Name", "Quantity"])
    data_df.to_excel('stock_control.xlsx', index=False)

def load_data():
    global data_df
    if os.path.exists('stock_control.xlsx'):
        data_df = pd.read_excel('stock_control.xlsx')
        for index, row in data_df.iterrows():
            list_data.append([row['Name'], row['Quantity']])

def add_product():
    product_name = name_entry.get()
    product_quantity = quantity_entry.get()

    if not product_name or not product_quantity:
        messagebox.showerror('Error', 'All fields are required')
        return

    list_data.append([product_name, product_quantity])
    save_data()
    display_data()

def remove_product():
    try:
        selected_product = list_box.curselection()[0]
        list_data.pop(selected_product)
        save_data()
        display_data()
    except IndexError:
        messagebox.showerror('Error', 'No product selected')

def display_data():
    list_box.delete(0, tk.END)
    for data in list_data:
        list_box.insert(tk.END, data)

def validar_entrada(valor):
    return valor.replace(',', '.').replace(',', '', 1).isdigit() or valor == ""

list_data = []
load_data()

window = tk.Tk()
window.title('Stock Control System')
window.geometry('400x400') # Definindo o tamanho do sistema
window.resizable(width=False, height=False) # Deixar o tamanho fixo sem resposividade

frame = tk.Frame(window)
frame.pack(pady=20)

name_label = tk.Label(frame, text='Nome do Produto')
name_label.grid(row=0, column=0)

name_entry = tk.Entry(frame)
name_entry.grid(row=0, column=1)

quantity_label = tk.Label(frame, text='Quantidade')
quantity_label.grid(row=1, column=0)

quantity_entry = tk.Entry(frame, validate="key", validatecommand=(window.register(validar_entrada), "%P"))
quantity_entry.grid(row=1, column=1)

add_button = tk.Button(frame, text='Adicionar produto', command=add_product)
add_button.grid(row=0, column=2, rowspan=2)

list_box = tk.Listbox(window, width=40, height=10)
list_box.pack(pady=20)

save_button = tk.Button(window, text='Salvar lista', command=save_data)
save_button.pack(padx=50)

remove_button = tk.Button(window, text='Remover Produto', command=remove_product)
remove_button.pack(pady=20)

display_data()

window.mainloop()